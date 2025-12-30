 
======================================== 
文件: 2-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 13:55:04 2025

@author: ASUS
"""

'''
（1）	编写程序，输入任意大的自然数，输出各位数字之和。
'''
x=int(input())
res=0
while x>0:
    res+=x%10
    x//=10
print(res) 
======================================== 
文件: 2-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 13:57:16 2025

@author: ASUS
"""

'''
（2）	编写程序，输入两个集合setA和setB，分别输出它们的交集、并集和差集setA-setB。
'''
setA=eval(input())
setB=eval(input())
print(setA&setB)
print(setA|setB)
print(setA-setB)
 
======================================== 
文件: 2-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:00:09 2025

@author: ASUS
"""

'''
（3）	编写程序，输入一个自然数，
输出它的二进制、八进制、十六进制表示形式。
'''
x=int(input())
print(bin(x))
print(oct(x))
print(hex(x)) 
======================================== 
文件: 2-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:02:28 2025

@author: ASUS
"""

'''
（4）	编写程序，输入一个包含若干整数的列表，
输出一个新列表，要求新列表中只包含原列表中的偶数。
'''
from random import randint
x=[ randint(0,10) for i in range(10) ]
res=[]
for i in range(len(x)):
    if x[i] % 2==0:
        res.append(x[i])
print(res) 
======================================== 
文件: 2-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:06:12 2025

@author: ASUS
"""

'''
（5）	编写程序，输入两个分别包含若干整数的列表lstA和lstB，
输出一个字典，要求使用列表lstA中的元素作为键，
列表lstB中的元素作为值，并且最终字典中的元素数量取决于lstA和lstB
中元素最少的列表的数量。
'''
lstA=eval(input())
lstB=eval(input())
lst_len_min=min(len(lstA),len(lstB))
res=dict()# or {}
#set() 集合
for i in range(lst_len_min):
    res[lstA[i]]=lstB[i]
print(res) 
======================================== 
文件: 2-6.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:21:16 2025

@author: ASUS
"""

'''
（6）	编写程序，输入一个包含若干整数的列表，
输出新列表，要求新列表中的所有元素来自于输入的列表，
并且降序排列。
'''
x=list(map(int,input().split(',')))
x.sort(reverse=True)
print(x) 
======================================== 
文件: 2-7.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:42:47 2025

@author: ASUS
"""

'''
（7）	编写程序，输入一个包含若干整数的列表，
输出列表中所有整数连乘的结果。
'''
from functools import reduce


lst=list(map(int,input().split(',')))
res=reduce(lambda x,y : x*y ,lst)
print(res)

 
======================================== 
文件: 2-8.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:56:09 2025

@author: ASUS
"""

'''
编写程序，输入两个各包含2个整数的列表，
分别表示城市中两个地点的坐标，输出两点之间的曼哈顿距离。
'''
x1,y1=list(map(int,input().split()))
x2,y2=list(map(int,input().split()))
print(abs(x1-x2)+abs(y1-y1)) 
======================================== 
文件: 2-9.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:59:32 2025

@author: ASUS
"""

'''
编写程序，输入包含若干集合的列表，输出这些集合的并集。
要求使用reduce()函数和lambda表达式完成。
'''
from functools import reduce
x=eval(input())
res=reduce(lambda x,y:x|y ,x)
print(res) 
======================================== 
文件: 3-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:06:54 2025

@author: ASUS
"""

'''
组合数求解
'''
def cni(n,i):
    minNI = min(i, n-i)
    result = 1
    res=1
    for j in range(0, minNI):
        result = result * (n-j) // (minNI-j)
        res = res * (n-j) // (j+1)
    return result,res
print(cni(5, 2)) 
======================================== 
文件: 3-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:39:37 2025

@author: ASUS
"""

'''
编写程序，输入一个大于2的自然数，
然后输出小于该数字的所有素数组成的列表。
(使用列表实现筛选法求素数)
'''
from math import sqrt
from sympy import isprime
def is_prime(a):
    if a<=1:
        return False
    else:
        for i in range(2,int(sqrt(a))+1):
            if a%i==0:
                return False
        return True

n=int(input())
res=[]
for i in range(2,n+1):
    if is_prime(i):
        res.append(i)
print(res)
 
======================================== 
文件: 3-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:48:08 2025

@author: ASUS
"""

'''
输入一个大于2的自然数，输出小于该数字的所有素数组成的集合。
（使用集合实现筛选法求素数）
'''
#差集
a=int(input())
res=set(range(a))
for i in range(2,int(a**0.5)+1):
    res-=set(range(i*i,a,i))
print(res)     
======================================== 
文件: 3-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 16:25:56 2025

@author: ASUS
"""

'''
使用filter()函数统计列表中所有非素数
'''
from sympy import isprime
from random import randint
a=[randint(0,100) for i in range(10)]
print(a)
b=list(filter(lambda x:not isprime(x) ,a))
print(b) 
======================================== 
文件: 3-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 16:45:57 2025

@author: ASUS
"""

'''
编写程序，输入等比数列的首项、
公比（不等于1且小于36的正整数）和一个自然数n，输出这个数列前n项的和。
关键步骤要求使用内置函数int()
'''
a=int(input())
q=int(input())
n=int(input())
res=0
if q==1 or q>=36:
    print("error")
else:
    res=a*(1-q**n)/(1-q)
    print(res) 
======================================== 
文件: 3-6.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 17:01:37 2025

@author: ASUS
"""

'''
编写程序，输入一个字符串，
输出其中出现次数最多的字符及其出现的次数。要求使用字典。
'''
x=input()
dic=dict()
for s in x:
    dic[s]=dic.get(s,0)+1
print(dic)
max_s=max(dic,key=dic.get)
print(max_s) 
======================================== 
文件: 4-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 11:25:47 2025

@author: phy12
"""
'''
编写程序，输入掷飞镖次数，然后输出圆周率近似值。
'''
from random import random
cnt=int(input())
res=0
x=[random() for i in range(cnt)]
y=[random() for i in range(cnt)]
for i in range(cnt):
    if x[i]**2+y[i]**2<=1:
        res+=1
print(res/cnt*4) 
======================================== 
文件: 4-2.py 
======================================== 
 
'''
对任意各位数字不相同的4位数，使用各位数字能组成的最大数减去能组成的最小数，
对得到的差重复这个操作，最终会得到的数字是什么？验证猜想。
'''
x=list(map(int,input()))
print(x)
def memo(x):
    x_1=sorted(x)#升序
    x_2=sorted(x,reverse=True)#降序
    res=(x_2[0]*1000+x_2[1]*100+x_2[2]*10+x_2[3])\
          -\
          (x_1[0]*1000+x_1[1]*100+x_1[2]*10+x_1[3])
    res=list(map(int,str(res)))
    return res
while True:
    print(memo(x))
    x=memo(x) 
======================================== 
文件: 4-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 14:55:11 2025

@author: ASUS
"""
'''
 编写程序，模拟决赛现场最终成绩计算过程。
 首先输入大于2的整数作为评委人数，然后依次输入每个评委的打分，
 要求每个分数都介于0和100之间。输入完所有评委打分之后，
 去掉一个最高分，去掉一个最低分，剩余分数的平均分即为该选手的最终得分。
'''
n=int(input())
score=list(map(int,input().split()))
if n<=2 or list(filter(lambda x:x>=100 ,score)):
    print("error")
else:
    score.sort()
    score=score[1:n-1]
    print(f"result score:{sum(score)/n:.2f}")
     
======================================== 
文件: 4-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 15:08:09 2025

@author: ASUS
"""

'''
 编写程序模拟猜数游戏。
 程序运行时，系统生成一个随机数，然后提示用户进行猜测，
 并根据用户输入进行必要的提示（猜对了、太大了、太小了），
 如果猜对则提前结束程序，如果次数用完仍没有猜对，
 提示游戏结束并给出正确答案。
'''
from random import randint
x=randint(0, 100)
cnt=5#次数
while cnt>0:
    g=int(input())
    cnt-=1
    if g>x:
        print("bigger")
    elif g<x:
        print("smaller")
    else:
        print("equal")
        break
if cnt==0 and g!=x:
    print("次数用尽") 
======================================== 
文件: 5-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 15:19:29 2025

@author: ASUS
"""

'''
假设一段楼梯共15个台阶，小朋友一步最多能上3个台阶。
编写程序计算小明上这段楼梯一共有多少种方法。
要求给出递推法和递归法两种代码。
'''
def floor1(x):
    if x==1 or x==0:
        return 1
    elif x==2:
        return 2
    else:
        return floor1(x-1)+floor1(x-2)+floor1(x-3)
print(floor1(15))

def floor2(f):
    if f==1:
        print(1)
    elif f==0:
        print(1)
    elif f==2:
        print(2)
    else:
        a,b,c=2,1,1
        while f>=3:
            a,b,c=a+b+c,a,b
            f-=1
        print(a)
floor2(15)
 
======================================== 
文件: 5-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:17:06 2025

@author: ASUS
"""

'''
编写程序，模拟抓土拨鼠小游戏。
假设一共有一排5个洞口，土拨鼠最开始的时候在其中一个洞口，
然后玩家随机打开一个洞口，如果里面有土拨鼠就抓到了。
如果洞口里没有土拨鼠就第二天再来抓，
但是第二天土拨鼠会在玩家来抓之前跳到隔壁洞口里。
如果在规定的次数内抓到了土拨鼠就提前结束游戏并提示成功；
如果规定的次数用完还没有抓到土拨鼠，就结束游戏并提示失败。
'''
from random import randint
hole=[0]*5
a=[-1,1]
choice=randint(0,4)
hole[choice]=1
cnt=5
while cnt>0:
    x=int(input())
    if x==choice:
        print("find it!")
        break
    elif choice==0:
        chioce=1
    elif choice==4:
        choice=3
    else:
        choice+=a[randint(0,1)]
    cnt-=1
if x!=choice:
    print("次数用尽") 
======================================== 
文件: 5-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:32:27 2025

@author: ASUS
"""

'''
两个玩家轮流从一堆物品中拿走一部分。
在每一步中，玩家可以自由选择拿走多少物品，
但是必须至少拿走一个并且最多只能拿走一半物品，
然后轮到下一个玩家。拿走最后一个物品的玩家输掉游戏。
编写程序，模拟该游戏。人机大战，机器使用最佳策略
'''
from random import randint, choice

def everyStep(n):
    half = n / 2
    m = 1
    # 所有可能满足条件的取法
    possible = []
    while True:
        rest = 2**m - 1
        if rest >= n:
            break
        if rest >= half:
            possible.append(n-rest)
        m = m+1
    # 如果至少存在一种取法使得剩余物品数量为2^n-1
    if possible:
        return choice(possible)
    # 无法使得剩余物品数量为2^n-1，随机取走一些
    return randint(1, int(half))

def smartNimuGame(n):
    while n > 1:
        # 人类玩家先走
        print("It's your turn, and we have {0} left.".format(n))
        # 确保人类玩家输入合法整数值
        while True:
            try:
                num = int(input('How many do you want to take:'))
                assert 1 <= num <= n//2
                break
            except:
                print('Must be between 1 and {0}'.format(n//2))
        n -= num
        if n == 1:
            return 'I fail.'
        # 计算机玩家拿走一些
        n -= everyStep(n)            
    else:
        return 'You fail.'

print(smartNimuGame(randint(1, 100)))
 
======================================== 
文件: 5-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:34:38 2025

@author: ASUS
"""
'''
消费者用力转动圆盘，然后圆盘慢慢停下来，
当圆盘恢复静止状态时，转盘上的指针所指的区域代表该消费者所中奖品。
如图所示：
假设共设一等奖（100元）、二等奖（20元）和三等奖（2元）这3个价值的奖品。
把圆盘从0到360度划分为3个区域，从0到30度对应一等奖，
30到108度对应二等奖，108到360度对应三等奖。
使用0到360之间的随机数表示消费者转动圆盘后指针所处位置。
'''
from random import randint
degree=randint(0,360)
if 0<=degree<30:
    print("一等奖")
elif 31<=degree<108:
    print("二等奖")
else :
    print("三等奖")
 
======================================== 
文件: 5-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:51:40 2025

@author: ASUS
"""

'''
假设你正参加一个有奖游戏节目，前方有3道门可以选择，
其中一个后面是汽车，另外两个后面是山羊。你选择一个门，
比如说1号门，主持人当然知道每个门后面是什么并且打开了另一个门，
比如说3号门，后面是一只山羊。这时，主持人会问你"你想改选2号门吗？"，
然后根据你的选择确定最终要打开的门，并确定你获得山羊（输）或者汽车（赢）。
编写程序，模拟上面的游戏，比如模拟20局，统计次数。
'''
from random import randint
from random import choice
def memo1():
    door=[0]*3
    door[randint(0, 2)]=1
    you_choice=randint(0,2)
    cnt=0
    op=choice([i for i in range(3) if i!=you_choice and door[i]!=1])
    #改选
    res1=[i for i in range(3) if i!=you_choice and i!=op][0]
    if door[res1]==1:
        cnt+=1
    return cnt
def memo2():
    door=[0]*3
    door[randint(0, 2)]=1
    you_choice=randint(0,2)
    cnt=0
    #op=choice([i for i in door if i!=you_choice and door[i]!=1])
    #不改选
    res2=you_choice
    if door[res2]==1:
        cnt+=1
    return cnt

cnt1,cnt2=0,0
test=100000
for _ in range(test):
    cnt1+=memo1()
    cnt2+=memo2()
print(cnt1/test,cnt2/test)















 
======================================== 
文件: 6-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 12:57:05 2025

@author: ASUS
"""

'''
把明文中每个英文字母替换为该字母在字母表中后面第k个字母，
如果后面第k个字符超出字母表的范围，则把字母表首尾相接，
也就是字母Z的下一个字母是A，字母z的下一个字母是a。
要求明文中的大写字母和小写字母分别进行处理，
大写字母加密后仍为大写字母，小写字母加密后仍为小写字母。
编写程序，输入一个字符串作为待加密的明文，
然后输入一个整数作为密钥，最后输出该字符串使用该密钥加密后的结果。
'''
mw=input()
n=int(input())
s=""
for i in mw:
    if 'a'<=i<='z':
        s+=chr((ord(i)-ord('a')+n)%26+ord('a'))
    if 'A'<=i<='Z':
        s+=chr((ord(i)-ord('A')+n)%26+ord('A'))
print(s) 
======================================== 
文件: 6-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:07:37 2025

@author: ASUS
"""

'''
把密码安全强度分为强密码、中高、中低、弱密码。
其中强密码表示字符串中同时含有数字、小写字母、大写字母、标点符号这4类字符，
而弱密码表示字符串中仅包含4类字符中的一种。编写程序，输入一个字符串，
输出该字符串作为密码时的安全强度。
'''
password=input()
c1,c2,c3,c4=0,0,0,0
for s in password:
    if '0'<=s<='9':
        c1=1
    if 'a'<=s<='z':
        c2=1
    if 'A'<=s<='Z':
        c3=1
    if s in ".,:;'":
        c4=1
res=sum([c1,c2,c3,c4])
if res==1:
    print("弱密码")
if res==2:
    print("中低密码")
if res==3:
    print("中高密码")
if res==4:
    print("强密码") 
======================================== 
文件: 6-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:26:45 2025

@author: ASUS
"""

'''
写程序，模拟打字练习程序的成绩评定。
假设origin为原始内容，userInput为用户练习时输入的内容，
要求用户输入的内容长度不能大于原始内容的长度。
如果对应位置的字符一致则认为正确，否则判定输入错误。
最后成绩为：正确的字符数量/原始字符串长度，按百分制输出，
要求保留2位小数。
'''
origin="Hello World"
userInput=input()
if len(userInput)>len(origin):
    print("请重新输入")
else:
    cnt=0
    for i in range(len(userInput)):
        if origin[i]==userInput[i]:
            cnt+=1
print(f"{round(cnt/len(origin)*100,2):.2f}","%")
    
     
======================================== 
文件: 6-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:41:11 2025

@author: ASUS
"""

'''
一般来说，在一封正常邮件中，
是不会出现太多类似于【、】、*、-、/这样的字符的。
如果一封邮件中包含的类似字符数量超过一定的比例，
我们可以直接认为是垃圾邮件，编写程序，对给定的邮件内容进行分类，
提示“垃圾邮件”或“正常邮件”。
'''
s="【】*-/"
str1="fshdfsdfsndsfsf/////////【【】【【】【--**】【】"
cnt=0
for i in str1:
    if i in s:
        cnt+=1
print(f"{round(cnt/len(str1)*100,2):.2f}","%") 
======================================== 
文件: 7-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:53:40 2025

@author: ASUS
"""

'''
编写一个程序demo.py，要求运行该程序后，生成demo_new.py文件，
其中内容与demo.py一样，只是在每一行的后面加上行号。要求行号以#开始，
并且所有行的#符号垂直对齐。
'''
filename="demo.py"
with open(filename,"r") as fp:
    lines=fp.readlines()
maxline=len(max(lines,key=len))
with open(filename[:-3]+"_new.py","w") as fp:
    lines=[line.rstrip().ljust(maxline)+"#"+str(i+1)+"\n" \
           for i,line in enumerate(lines)]
    fp.writelines(lines) 
======================================== 
文件: 7-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 14:28:32 2025

@author: ASUS
"""

'''
2， 编写程序，实现磁盘垃圾文件清理功能。
要求指定要清理的文件夹，
然后删除该文件夹及其子文件夹中所有扩展名为
tmp、log、obj、txt以及大小为0的文件。
函数参数为文件夹名称及路径。
'''
from os.path import isdir, join, splitext, getsize
from os import remove, listdir


filetypes=[".tmp",".log",".obj",".txt"]
def delfiles(pwd):
    if not isdir(pwd):
        return 
    for filename in listdir(pwd):
        temp=join(pwd,filename)
        if isdir(temp):
            delfiles(temp)
        if splitext(temp)[1] in filetypes or getsize(temp)==0:
            remove(temp)
pwd=r"D:\testDirectory"
delfiles(pwd)
 
======================================== 
文件: 7-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 16:15:56 2025

@author: ASUS
"""

'''
安装Python扩展库python-docx，然后读取一个Word文章中所有段落的文本，
查找并输出其中所有AABB形式的词语，例如踏踏实实、密密麻麻、简简单单、时时刻刻。
'''
import re
from docx import Document
doc=Document("text.docx")
for para in doc.paragraphs:
    text=para.text
    pattern = r'(([\u4e00-\u9fa5])\2([\u4e00-\u9fa5])\3)'
    r = re.findall(pattern, text)
    if r:
        for word in r:
            print(word[0]) 
======================================== 
文件: 7-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 16:34:55 2025

@author: ASUS
"""

'''
编写程序，读取Word文件中的所有段落文本，
然后输出其中所有红色的文本和加粗的文本以及同时具有这两种属性的文本。
'''
from docx import Document
from docx.shared import RGBColor

boldText = []
redText = []
doc = Document('test.docx')
for p in doc.paragraphs:
    for r in p.runs:
        # 加粗字体
        if r.bold:
            boldText.append(r.text)
        # 红色字体
        if r.font.color.rgb == RGBColor(255,0,0):
            redText.append(r.text)

result = {'red text': redText,
          'bold text': boldText,
          'both': set(redText) & set(boldText)}
#  输出结果
for title in result.keys():
    print(title.center(30, '='))
    for text in result[title]:
        print(text)
 
======================================== 
文件: merged.py 
======================================== 
 
 
======================================== 
文件: 2-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 13:55:04 2025

@author: ASUS
"""

'''
（1）	编写程序，输入任意大的自然数，输出各位数字之和。
'''
x=int(input())
res=0
while x>0:
    res+=x%10
    x//=10
print(res) 
======================================== 
文件: 2-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 13:57:16 2025

@author: ASUS
"""

'''
（2）	编写程序，输入两个集合setA和setB，分别输出它们的交集、并集和差集setA-setB。
'''
setA=eval(input())
setB=eval(input())
print(setA&setB)
print(setA|setB)
print(setA-setB)
 
======================================== 
文件: 2-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:00:09 2025

@author: ASUS
"""

'''
（3）	编写程序，输入一个自然数，
输出它的二进制、八进制、十六进制表示形式。
'''
x=int(input())
print(bin(x))
print(oct(x))
print(hex(x)) 
======================================== 
文件: 2-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:02:28 2025

@author: ASUS
"""

'''
（4）	编写程序，输入一个包含若干整数的列表，
输出一个新列表，要求新列表中只包含原列表中的偶数。
'''
from random import randint
x=[ randint(0,10) for i in range(10) ]
res=[]
for i in range(len(x)):
    if x[i] % 2==0:
        res.append(x[i])
print(res) 
======================================== 
文件: 2-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:06:12 2025

@author: ASUS
"""

'''
（5）	编写程序，输入两个分别包含若干整数的列表lstA和lstB，
输出一个字典，要求使用列表lstA中的元素作为键，
列表lstB中的元素作为值，并且最终字典中的元素数量取决于lstA和lstB
中元素最少的列表的数量。
'''
lstA=eval(input())
lstB=eval(input())
lst_len_min=min(len(lstA),len(lstB))
res=dict()# or {}
#set() 集合
for i in range(lst_len_min):
    res[lstA[i]]=lstB[i]
print(res) 
======================================== 
文件: 2-6.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:21:16 2025

@author: ASUS
"""

'''
（6）	编写程序，输入一个包含若干整数的列表，
输出新列表，要求新列表中的所有元素来自于输入的列表，
并且降序排列。
'''
x=list(map(int,input().split(',')))
x.sort(reverse=True)
print(x) 
======================================== 
文件: 2-7.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:42:47 2025

@author: ASUS
"""

'''
（7）	编写程序，输入一个包含若干整数的列表，
输出列表中所有整数连乘的结果。
'''
from functools import reduce


lst=list(map(int,input().split(',')))
res=reduce(lambda x,y : x*y ,lst)
print(res)

 
======================================== 
文件: 2-8.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:56:09 2025

@author: ASUS
"""

'''
编写程序，输入两个各包含2个整数的列表，
分别表示城市中两个地点的坐标，输出两点之间的曼哈顿距离。
'''
x1,y1=list(map(int,input().split()))
x2,y2=list(map(int,input().split()))
print(abs(x1-x2)+abs(y1-y1)) 
======================================== 
文件: 2-9.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 14:59:32 2025

@author: ASUS
"""

'''
编写程序，输入包含若干集合的列表，输出这些集合的并集。
要求使用reduce()函数和lambda表达式完成。
'''
from functools import reduce
x=eval(input())
res=reduce(lambda x,y:x|y ,x)
print(res) 
======================================== 
文件: 3-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:06:54 2025

@author: ASUS
"""

'''
组合数求解
'''
def cni(n,i):
    minNI = min(i, n-i)
    result = 1
    res=1
    for j in range(0, minNI):
        result = result * (n-j) // (minNI-j)
        res = res * (n-j) // (j+1)
    return result,res
print(cni(5, 2)) 
======================================== 
文件: 3-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:39:37 2025

@author: ASUS
"""

'''
编写程序，输入一个大于2的自然数，
然后输出小于该数字的所有素数组成的列表。
(使用列表实现筛选法求素数)
'''
from math import sqrt
from sympy import isprime
def is_prime(a):
    if a<=1:
        return False
    else:
        for i in range(2,int(sqrt(a))+1):
            if a%i==0:
                return False
        return True

n=int(input())
res=[]
for i in range(2,n+1):
    if is_prime(i):
        res.append(i)
print(res)
 
======================================== 
文件: 3-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 15:48:08 2025

@author: ASUS
"""

'''
输入一个大于2的自然数，输出小于该数字的所有素数组成的集合。
（使用集合实现筛选法求素数）
'''
#差集
a=int(input())
res=set(range(a))
for i in range(2,int(a**0.5)+1):
    res-=set(range(i*i,a,i))
print(res)     
======================================== 
文件: 3-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 16:25:56 2025

@author: ASUS
"""

'''
使用filter()函数统计列表中所有非素数
'''
from sympy import isprime
from random import randint
a=[randint(0,100) for i in range(10)]
print(a)
b=list(filter(lambda x:not isprime(x) ,a))
print(b) 
======================================== 
文件: 3-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 16:45:57 2025

@author: ASUS
"""

'''
编写程序，输入等比数列的首项、
公比（不等于1且小于36的正整数）和一个自然数n，输出这个数列前n项的和。
关键步骤要求使用内置函数int()
'''
a=int(input())
q=int(input())
n=int(input())
res=0
if q==1 or q>=36:
    print("error")
else:
    res=a*(1-q**n)/(1-q)
    print(res) 
======================================== 
文件: 3-6.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Tue Dec 23 17:01:37 2025

@author: ASUS
"""

'''
编写程序，输入一个字符串，
输出其中出现次数最多的字符及其出现的次数。要求使用字典。
'''
x=input()
dic=dict()
for s in x:
    dic[s]=dic.get(s,0)+1
print(dic)
max_s=max(dic,key=dic.get)
print(max_s) 
======================================== 
文件: 4-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 11:25:47 2025

@author: phy12
"""
'''
编写程序，输入掷飞镖次数，然后输出圆周率近似值。
'''
from random import random
cnt=int(input())
res=0
x=[random() for i in range(cnt)]
y=[random() for i in range(cnt)]
for i in range(cnt):
    if x[i]**2+y[i]**2<=1:
        res+=1
print(res/cnt*4) 
======================================== 
文件: 4-2.py 
======================================== 
 
'''
对任意各位数字不相同的4位数，使用各位数字能组成的最大数减去能组成的最小数，
对得到的差重复这个操作，最终会得到的数字是什么？验证猜想。
'''
x=list(map(int,input()))
print(x)
def memo(x):
    x_1=sorted(x)#升序
    x_2=sorted(x,reverse=True)#降序
    res=(x_2[0]*1000+x_2[1]*100+x_2[2]*10+x_2[3])\
          -\
          (x_1[0]*1000+x_1[1]*100+x_1[2]*10+x_1[3])
    res=list(map(int,str(res)))
    return res
while True:
    print(memo(x))
    x=memo(x) 
======================================== 
文件: 4-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 14:55:11 2025

@author: ASUS
"""
'''
 编写程序，模拟决赛现场最终成绩计算过程。
 首先输入大于2的整数作为评委人数，然后依次输入每个评委的打分，
 要求每个分数都介于0和100之间。输入完所有评委打分之后，
 去掉一个最高分，去掉一个最低分，剩余分数的平均分即为该选手的最终得分。
'''
n=int(input())
score=list(map(int,input().split()))
if n<=2 or list(filter(lambda x:x>=100 ,score)):
    print("error")
else:
    score.sort()
    score=score[1:n-1]
    print(f"result score:{sum(score)/n:.2f}")
     
======================================== 
文件: 4-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 15:08:09 2025

@author: ASUS
"""

'''
 编写程序模拟猜数游戏。
 程序运行时，系统生成一个随机数，然后提示用户进行猜测，
 并根据用户输入进行必要的提示（猜对了、太大了、太小了），
 如果猜对则提前结束程序，如果次数用完仍没有猜对，
 提示游戏结束并给出正确答案。
'''
from random import randint
x=randint(0, 100)
cnt=5#次数
while cnt>0:
    g=int(input())
    cnt-=1
    if g>x:
        print("bigger")
    elif g<x:
        print("smaller")
    else:
        print("equal")
        break
if cnt==0 and g!=x:
    print("次数用尽") 
======================================== 
文件: 5-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 15:19:29 2025

@author: ASUS
"""

'''
假设一段楼梯共15个台阶，小朋友一步最多能上3个台阶。
编写程序计算小明上这段楼梯一共有多少种方法。
要求给出递推法和递归法两种代码。
'''
def floor1(x):
    if x==1 or x==0:
        return 1
    elif x==2:
        return 2
    else:
        return floor1(x-1)+floor1(x-2)+floor1(x-3)
print(floor1(15))

def floor2(f):
    if f==1:
        print(1)
    elif f==0:
        print(1)
    elif f==2:
        print(2)
    else:
        a,b,c=2,1,1
        while f>=3:
            a,b,c=a+b+c,a,b
            f-=1
        print(a)
floor2(15)
 
======================================== 
文件: 5-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:17:06 2025

@author: ASUS
"""

'''
编写程序，模拟抓土拨鼠小游戏。
假设一共有一排5个洞口，土拨鼠最开始的时候在其中一个洞口，
然后玩家随机打开一个洞口，如果里面有土拨鼠就抓到了。
如果洞口里没有土拨鼠就第二天再来抓，
但是第二天土拨鼠会在玩家来抓之前跳到隔壁洞口里。
如果在规定的次数内抓到了土拨鼠就提前结束游戏并提示成功；
如果规定的次数用完还没有抓到土拨鼠，就结束游戏并提示失败。
'''
from random import randint
hole=[0]*5
a=[-1,1]
choice=randint(0,4)
hole[choice]=1
cnt=5
while cnt>0:
    x=int(input())
    if x==choice:
        print("find it!")
        break
    elif choice==0:
        chioce=1
    elif choice==4:
        choice=3
    else:
        choice+=a[randint(0,1)]
    cnt-=1
if x!=choice:
    print("次数用尽") 
======================================== 
文件: 5-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:32:27 2025

@author: ASUS
"""

'''
两个玩家轮流从一堆物品中拿走一部分。
在每一步中，玩家可以自由选择拿走多少物品，
但是必须至少拿走一个并且最多只能拿走一半物品，
然后轮到下一个玩家。拿走最后一个物品的玩家输掉游戏。
编写程序，模拟该游戏。人机大战，机器使用最佳策略
'''
from random import randint, choice

def everyStep(n):
    half = n / 2
    m = 1
    # 所有可能满足条件的取法
    possible = []
    while True:
        rest = 2**m - 1
        if rest >= n:
            break
        if rest >= half:
            possible.append(n-rest)
        m = m+1
    # 如果至少存在一种取法使得剩余物品数量为2^n-1
    if possible:
        return choice(possible)
    # 无法使得剩余物品数量为2^n-1，随机取走一些
    return randint(1, int(half))

def smartNimuGame(n):
    while n > 1:
        # 人类玩家先走
        print("It's your turn, and we have {0} left.".format(n))
        # 确保人类玩家输入合法整数值
        while True:
            try:
                num = int(input('How many do you want to take:'))
                assert 1 <= num <= n//2
                break
            except:
                print('Must be between 1 and {0}'.format(n//2))
        n -= num
        if n == 1:
            return 'I fail.'
        # 计算机玩家拿走一些
        n -= everyStep(n)            
    else:
        return 'You fail.'

print(smartNimuGame(randint(1, 100)))
 
======================================== 
文件: 5-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:34:38 2025

@author: ASUS
"""
'''
消费者用力转动圆盘，然后圆盘慢慢停下来，
当圆盘恢复静止状态时，转盘上的指针所指的区域代表该消费者所中奖品。
如图所示：
假设共设一等奖（100元）、二等奖（20元）和三等奖（2元）这3个价值的奖品。
把圆盘从0到360度划分为3个区域，从0到30度对应一等奖，
30到108度对应二等奖，108到360度对应三等奖。
使用0到360之间的随机数表示消费者转动圆盘后指针所处位置。
'''
from random import randint
degree=randint(0,360)
if 0<=degree<30:
    print("一等奖")
elif 31<=degree<108:
    print("二等奖")
else :
    print("三等奖")
 
======================================== 
文件: 5-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Wed Dec 24 16:51:40 2025

@author: ASUS
"""

'''
假设你正参加一个有奖游戏节目，前方有3道门可以选择，
其中一个后面是汽车，另外两个后面是山羊。你选择一个门，
比如说1号门，主持人当然知道每个门后面是什么并且打开了另一个门，
比如说3号门，后面是一只山羊。这时，主持人会问你"你想改选2号门吗？"，
然后根据你的选择确定最终要打开的门，并确定你获得山羊（输）或者汽车（赢）。
编写程序，模拟上面的游戏，比如模拟20局，统计次数。
'''
from random import randint
from random import choice
def memo1():
    door=[0]*3
    door[randint(0, 2)]=1
    you_choice=randint(0,2)
    cnt=0
    op=choice([i for i in range(3) if i!=you_choice and door[i]!=1])
    #改选
    res1=[i for i in range(3) if i!=you_choice and i!=op][0]
    if door[res1]==1:
        cnt+=1
    return cnt
def memo2():
    door=[0]*3
    door[randint(0, 2)]=1
    you_choice=randint(0,2)
    cnt=0
    #op=choice([i for i in door if i!=you_choice and door[i]!=1])
    #不改选
    res2=you_choice
    if door[res2]==1:
        cnt+=1
    return cnt

cnt1,cnt2=0,0
test=100000
for _ in range(test):
    cnt1+=memo1()
    cnt2+=memo2()
print(cnt1/test,cnt2/test)















 
======================================== 
文件: 6-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 12:57:05 2025

@author: ASUS
"""

'''
把明文中每个英文字母替换为该字母在字母表中后面第k个字母，
如果后面第k个字符超出字母表的范围，则把字母表首尾相接，
也就是字母Z的下一个字母是A，字母z的下一个字母是a。
要求明文中的大写字母和小写字母分别进行处理，
大写字母加密后仍为大写字母，小写字母加密后仍为小写字母。
编写程序，输入一个字符串作为待加密的明文，
然后输入一个整数作为密钥，最后输出该字符串使用该密钥加密后的结果。
'''
mw=input()
n=int(input())
s=""
for i in mw:
    if 'a'<=i<='z':
        s+=chr((ord(i)-ord('a')+n)%26+ord('a'))
    if 'A'<=i<='Z':
        s+=chr((ord(i)-ord('A')+n)%26+ord('A'))
print(s) 
======================================== 
文件: 6-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:07:37 2025

@author: ASUS
"""

'''
把密码安全强度分为强密码、中高、中低、弱密码。
其中强密码表示字符串中同时含有数字、小写字母、大写字母、标点符号这4类字符，
而弱密码表示字符串中仅包含4类字符中的一种。编写程序，输入一个字符串，
输出该字符串作为密码时的安全强度。
'''
password=input()
c1,c2,c3,c4=0,0,0,0
for s in password:
    if '0'<=s<='9':
        c1=1
    if 'a'<=s<='z':
        c2=1
    if 'A'<=s<='Z':
        c3=1
    if s in ".,:;'":
        c4=1
res=sum([c1,c2,c3,c4])
if res==1:
    print("弱密码")
if res==2:
    print("中低密码")
if res==3:
    print("中高密码")
if res==4:
    print("强密码") 
======================================== 
文件: 6-3.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:26:45 2025

@author: ASUS
"""

'''
写程序，模拟打字练习程序的成绩评定。
假设origin为原始内容，userInput为用户练习时输入的内容，
要求用户输入的内容长度不能大于原始内容的长度。
如果对应位置的字符一致则认为正确，否则判定输入错误。
最后成绩为：正确的字符数量/原始字符串长度，按百分制输出，
要求保留2位小数。
'''
origin="Hello World"
userInput=input()
if len(userInput)>len(origin):
    print("请重新输入")
else:
    cnt=0
    for i in range(len(userInput)):
        if origin[i]==userInput[i]:
            cnt+=1
print(f"{round(cnt/len(origin)*100,2):.2f}","%")
    
     
======================================== 
文件: 6-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:41:11 2025

@author: ASUS
"""

'''
一般来说，在一封正常邮件中，
是不会出现太多类似于【、】、*、-、/这样的字符的。
如果一封邮件中包含的类似字符数量超过一定的比例，
我们可以直接认为是垃圾邮件，编写程序，对给定的邮件内容进行分类，
提示“垃圾邮件”或“正常邮件”。
'''
s="【】*-/"
str1="fshdfsdfsndsfsf/////////【【】【【】【--**】【】"
cnt=0
for i in str1:
    if i in s:
        cnt+=1
print(f"{round(cnt/len(str1)*100,2):.2f}","%") 
======================================== 
文件: 7-1.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 13:53:40 2025

@author: ASUS
"""

'''
编写一个程序demo.py，要求运行该程序后，生成demo_new.py文件，
其中内容与demo.py一样，只是在每一行的后面加上行号。要求行号以#开始，
并且所有行的#符号垂直对齐。
'''
filename="demo.py"
with open(filename,"r") as fp:
    lines=fp.readlines()
maxline=len(max(lines,key=len))
with open(filename[:-3]+"_new.py","w") as fp:
    lines=[line.rstrip().ljust(maxline)+"#"+str(i+1)+"\n" \
           for i,line in enumerate(lines)]
    fp.writelines(lines) 
======================================== 
文件: 7-2.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 14:28:32 2025

@author: ASUS
"""

'''
2， 编写程序，实现磁盘垃圾文件清理功能。
要求指定要清理的文件夹，
然后删除该文件夹及其子文件夹中所有扩展名为
tmp、log、obj、txt以及大小为0的文件。
函数参数为文件夹名称及路径。
'''
from os.path import isdir, join, splitext, getsize
from os import remove, listdir


filetypes=[".tmp",".log",".obj",".txt"]
def delfiles(pwd):
    if not isdir(pwd):
        return 
    for filename in listdir(pwd):
        temp=join(pwd,filename)
        if isdir(temp):
            delfiles(temp)
        if splitext(temp)[1] in filetypes or getsize(temp)==0:
            remove(temp)
pwd=r"D:\testDirectory"
delfiles(pwd)
 
======================================== 
文件: 7-4.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 16:15:56 2025

@author: ASUS
"""

'''
安装Python扩展库python-docx，然后读取一个Word文章中所有段落的文本，
查找并输出其中所有AABB形式的词语，例如踏踏实实、密密麻麻、简简单单、时时刻刻。
'''
import re
from docx import Document
doc=Document("text.docx")
for para in doc.paragraphs:
    text=para.text
    pattern = r'(([\u4e00-\u9fa5])\2([\u4e00-\u9fa5])\3)'
    r = re.findall(pattern, text)
    if r:
        for word in r:
            print(word[0]) 
======================================== 
文件: 7-5.py 
======================================== 
 
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 28 16:34:55 2025

@author: ASUS
"""

'''
编写程序，读取Word文件中的所有段落文本，
然后输出其中所有红色的文本和加粗的文本以及同时具有这两种属性的文本。
'''
from docx import Document
from docx.shared import RGBColor

boldText = []
redText = []
doc = Document('test.docx')
for p in doc.paragraphs:
    for r in p.runs:
        # 加粗字体
        if r.bold:
            boldText.append(r.text)
        # 红色字体
        if r.font.color.rgb == RGBColor(255,0,0):
            redText.append(r.text)

result = {'red text': redText,
          'bold text': boldText,
          'both': set(redText) & set(boldText)}
#  输出结果
for title in result.keys():
    print(title.center(30, '='))
    for text in result[title]:
        print(text)
 
